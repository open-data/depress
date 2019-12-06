from pathlib import Path

# noinspection PyPackageRequirements
from docx import Document
from docx.shared import Cm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from deplane.md_to_docx import insert_markdown


def write_docx(schema, filename, trans, lang):
    _ = trans.gettext

    document = Document(Path(__file__).parent / 'ENG_2_Colour.docx')
    section = document.sections[0]

    title = schema['title'][lang]

    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].text = ''
    section.header.paragraphs[0].text = _('Template Guide: {title}').format(title=title)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # replace existing "TITLE" / "Subtitle" text
    document.paragraphs[0].runs[0].font.size = Cm(1.2) # template title too large
    document.paragraphs[0].runs[0].text = '\n\n\n' + _('Data Element Profile')
    document.paragraphs[0].runs[2].text = '\n' + title
    # clear out the rest of the example text
    for para in reversed(document.paragraphs[2:]):
        delete_paragraph(para)

    document.add_heading(_('Legend'), 1)
    document.add_paragraph(_(
        'The following sample table provides a description of each field you will see '
        'for all contract elements:'
    ))

    table = document.add_table(rows=0, cols=2)
    table.autofit = False
    def trow(att, desc):
        cells = table.add_row().cells
        cells[0].text = att
        cells[1].text = desc

    trow(_('Attribute'), _('Attribute Description'))
    trow(
        _('Field Name EN'),
        _('This text should correspond directly with the field name in your template in English'),
    )
    trow(
        _('Field Name FR'),
        _('This text should correspond directly with the field name in your template in French'),
    )
    trow(
        _('Description EN'),
        _('This provides a brief description of the element in English'),
    )
    trow(
        _('Description FR'),
        _('This provides a brief description of the element in French'),
    )
    cells = table.add_row().cells
    cells[0].text = _('Obligation')
    delete_paragraph(cells[1].paragraphs[0])
    insert_markdown(
        cells[1],
        _('''
Indicates whether the element is required to always or sometimes be present 
(i.e., contain a value). Options are:

- Mandatory
- Mandatory, conditional
- Optional'''),
    )
    trow(
        _('Condition'),
        _('Describes the condition or conditions according to which a value shall be present'),
    )
    cells = table.add_row().cells
    cells[0].text = _('Format Type')
    delete_paragraph(cells[1].paragraphs[0])
    insert_markdown(
        cells[1],
        _('''
Indicates the required format of the values, if any, at the file level.
“Free text” indicates that the value may be input using natural language
(i.e., there is no constraint) while “single choice” or “multiple choice”
indicates the values are restricted to a controlled list.)

Controlled List Values:

Code | English | French
--- | --- | ---
CODE1 | English Description 1 | French Description 1
CODE2 | English Description 2 | French Description 2''')
    )
    trow(
        _('Validation'),
        _('Indicates what the system will accept in this field'),
    )
    trow(
        _('Validation Errors'),
        _(
            'This section indicates when an error has been made. '
            'It will detail the error and provide instruction on how to correct it.'
        ),
    )
    trow(
        _('Example Value'),
        _(
            'Provide one or more real examples of the values that may appear, '
            'e.g. “CODE1” or “Family Services Reform Program”'
        ),
    )

    format_table(
        table,
        widths=[Cm(4.69), Cm(11.80)],
        top_color='d9d9d9',
        left_color='c6d9f1',
    )

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)


    document.add_page_break()

    document.save(filename)


def format_table(table, widths, top_color=None, left_color=None):
    # for Word
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w
    # for Libreoffice
    for i, w in enumerate(widths):
        table.columns[i].width = w

    if left_color:
        for cell in table.columns[0].cells:
            element = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), left_color))
            cell._tc.get_or_add_tcPr().append(element)
    if top_color:
        for cell in table.rows[0].cells:
            element = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), top_color))
            cell._tc.get_or_add_tcPr().append(element)


def delete_paragraph(para):
    p = para._element
    p.getparent().remove(p)
    p._p = p._element = None